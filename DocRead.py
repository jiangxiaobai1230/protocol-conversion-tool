# -*- coding: utf-8 -*-
"""
Created on Thu May 30 12:13:42 2024

@author: jinn
"""
from docx import Document
import ArrayCmp
import pandas as pd
import numpy as np
import os
import filemanage as fm
from docx.oxml.shared import OxmlElement,qn

def get_doc_table_row_context(table,array,rowIndex,clearHead = True,unique = True):
    """读取表格中的一行数据放在array里面，rowindex是表格行号
    如果clearHead为真，则去掉读取的第一个元素
    如果unique是真则去掉重复的元素
    """
    array.clear()
    tmpArray = []
    
    if True == clearHead:
        startIndex = 1
    else:
        startIndex = 0
            
    arraylen = len(table.columns)
    for i in range(startIndex,arraylen):
       # cellTmp = table.cell(rowIndex,i)
       # if cellTmp.merge_vertical or cellTmp.horizontal
        tmpText = table.cell(rowIndex,i).text
        #去除换行符
        cleanText = tmpText.replace('\n', '')
        tmpArray.append(table.cell(rowIndex,i).text)
    if True == unique :       
        "去掉合并单元格的影响"
        ArrayCmp.get_unique_array(tmpArray,array)
    else:
        ArrayCmp.list_copy(tmpArray,array)
        
def get_doc_tabel_row_appoint_context(table,array,rowIndex,indexArray):
    "根据indexArray读取表格内容，rowIndex代表行号"
    array.clear()
    tmpArray = []
    tmpAppoint = []
    get_doc_table_row_context(table, tmpArray, rowIndex,False,False)
    tmpAppoint = ArrayCmp.get_appoint_array(tmpArray, indexArray)
    ArrayCmp.list_copy(tmpAppoint, array)
    
def get_doc_table_head_index(table,array,rowIndex,indexArray):
    """
    获取信息表的表头，返回表头各元素的列索引，通过该索引可以找到对应的元素

    Parameters
    ----------
    table : TYPE
        表格句柄.
    array : TYPE
        初始化用来存放表头数据内容的数组.
    rowIndex : TYPE
        行号.

    Returns
    -------
    表头各元素索引

    """
    indexArray.clear()
    array.clear()    
    tmpArray = []
    tmpIndex = []
    get_doc_table_row_context(table, tmpArray, rowIndex,False,False)
    tmpIndex = ArrayCmp.get_unique_array(tmpArray, array)
    ArrayCmp.list_copy(tmpIndex, indexArray)
  
    
def get_doc_table_col_context(table,colIndex,clearHead = True):
    """读取一列数据
    如果clearHead为真，则去掉读取的第一个元素
    """
    array = []
    if True == clearHead:
        startIndex = 1
    else:
        startIndex = 0
    arraylen = len(table.rows)
    for i in range(startIndex,arraylen):
        array.append(table.cell(i,colIndex).text) 
    return array


def mate_table_head(headOriArray,baseline):
    """
    通过匹配数组匹配当前表头，如果一直则认为找到表格
    如果有合并单元格情况，则进行合并，维数应与baseline一致

    Parameters
    ----------

    baseline : TYPE
      表头匹配数组.

    Returns
    -------
    是否与匹配数组一致.

    """

    headArray = []
    headOldArray = []
    #表头均为首行
    ArrayCmp.get_unique_array(headOriArray, headOldArray)
    #获取仅含有中文的表头
    ArrayCmp.get_clean_chsArray(headOldArray, headArray)

    tableColNum = len(headArray)
    baselineNum = len(baseline)
    if tableColNum == baselineNum:
       
        return ArrayCmp.are_lists_equal(baseline,headOriArray)
    else:
        #存在合并单元格
        print("mate_table_head:维数不相等：","baseline:",baseline,"headOriArray",headOriArray)
        return False

def mate_table_head_vague(rowArray,baseline):
    "仅检查rowArray中是否包含baseline中的所有元素"
    mateNum = len(baseline)
    mateResult = 0
    for mateWord in baseline:
        if mateWord in rowArray:
            mateResult += 1
    
    if mateResult == mateNum:
        return True
    else:
        return False
def check_unique_valueindex(nums):
    # 使用 set 来存储已经出现的数字
    seen = set()
    # 使用列表来收集重复的数字索引
    unique = []
    for i, num in enumerate(nums):
        # 如果数字已经在seen中，则是重复的，收集其索引
        if num in seen:
            continue
        else:
            seen.add(num)
            unique.append(i)
    return unique
 

def get_msg_read_index(table,colIndex):
    """
    判断table[所有行][colIndex列]的数值，判断数值的合并情况
    选出唯一的数值索引
    用于表格行合并情况下的信息提取
    Parameters
    ----------
    table : TYPE
        DESCRIPTION.
    colIndex : TYPE
        用于判断合并情况的列索引.
    indexArray : TYPE
        挑出合并行的列索引.

    Returns
    -------
    None.

    """
    msgNameArray = get_doc_table_col_context(table, colIndex,False)
    #如果第一个元素和后面的重合则警告
    if msgNameArray[0] in msgNameArray[1:len(msgNameArray)]:
        print("get_msg_read_index:是否标题与内容重合",msgNameArray)
    #先判断是否有重复值
    if len(msgNameArray) != len(set(msgNameArray)):
        #有重复数
        indexArray = check_unique_valueindex(msgNameArray)
    else:
        indexArray = list(range(0,len(msgNameArray)))
    return indexArray

class DataSet:
    "用于存放消息下数据字段的各类属性"
    """msgNameArray = pd.DataFrame(columns = ["名称","信源系统码","信源机器码",
                                           "信宿系统码","信宿机器码","子地址",
                                           "数据段长度","ID"])
    data = pd.DataFrame(columns = ["内容","子内容","类型（bit）","转换类型",
                                   "判读公式","转换公式","单位","备注"])
    """
    msgAddrArray = []
    msgNameArray = []
    msgIDArray = []
    busdata = []
    def _init_(self):
        pass
    
    
def get_bus_msg(doc,filename,outputpath):
    """匹配总线信息表头，如果是对的表头则提取数据放入dataframe"""
    "遍历文档表格，直至找到第一个符合表头要求的表格"
    
    busbaseline = ["信息名称","信息标识"]
    """
    msgbaseline = ["序号",	"内容",	"类型",	"值域",	"单位",	
                   "数据处理方法"]
    """
    #临时调整模板
    msgbaseline = ["序号",	"内容","类型","值域",	"单位",	
                   "数据处理方法"]
    #用于挑出数据的索引
    templateName = ["内容","子内容","类型（bit）","转换类型",
                                "判读公式","转换公式","单位","备注"]
    #按该数组映射： 内容，子内容（空），类型，转换类型，公式（空），转换公式（空），单位，数据处理方法
    #里面的索引值是msgbaseline匹配到template的索引值
    msgMapIndex = pd.Series(data = [1,"",2,"","","",4,5],index = templateName)

    "名称"
    msgNameArray = []
    "地址描述字符串数组，如bc-rt3sa2L5,再经过函数拆分出关键词"
    msgAddrArray = []
    "id暂时为空"
    msgIDArray = []
    rowArray = []
    "数据字段的表头索引，防止存在列表合并单元导致取错数据"
    indexArray = []
    dataStartIndex = 5

    busdata = pd.DataFrame(columns = ["内容","子内容","类型（bit）","转换类型",
                                "判读公式","转换公式","单位","备注"])
    index = 0
    
    for tmpTable in doc.tables:
        "匹配表头"
        get_doc_table_row_context(tmpTable, rowArray, 0,False)
        if mate_table_head_vague(rowArray, busbaseline):
            "找到了表格，直接取数"
            "取信息名称，位置固定"
            tmpmsgname = rowArray[1]
            #msgNameArray.append(rowArray[1])
            "取信源信宿字符串"
            get_doc_table_row_context(tmpTable, rowArray, 1,False)
            tmpmsgaddr = rowArray[1]
            #msgAddrArray.append(rowArray[1])
            "先匹配数据段表头再取数,不去重，获得索引"
            
            get_doc_table_head_index(tmpTable, rowArray, 4,indexArray)
            if mate_table_head(rowArray, msgbaseline):
                msgNameArray.append(tmpmsgname)
                msgAddrArray.append(tmpmsgaddr)
                #id暂时赋值空
                msgIDArray.append("")
                #获取数据项列索引，去掉重复的行,采用内容进行查重
                uniq_index = get_msg_read_index(tmpTable, 1)
                
                "数据"
                tableRowNum = len(uniq_index)
                iterNum = dataStartIndex

                while iterNum < tableRowNum:   
                     "取内容,类型当全部写完后批量转换,防止合并单元格影响取数采用表头的索引"
                     get_doc_tabel_row_appoint_context(tmpTable,rowArray, uniq_index[iterNum], indexArray)
                     "去掉注"
                     "如果不小于0则是有注"
                     if rowArray[0].find('注') >= 0:
                         pass
                     #如果是最后一行且含有CRC则认为是校验码，删除
                     elif rowArray[msgMapIndex.loc["内容"]].find('CRC') >=0 and tableRowNum-iterNum == 1:
                         pass
                     else:
                         if len(rowArray) != len(msgbaseline):
                             print("get_bus_msg:总线数据列错误")
                         else:
                             "转换类型"
                             (bitType,newType) = ArrayCmp.standard_datatype(rowArray[msgMapIndex.loc["类型（bit）"]])
                             """
                             #内容，子内容（空），类型，转换类型，公式（空），转换公式（空），单位，数据处理方法
                             busdata.loc[index] = [rowArray[1],"",bitType,newType,
                                                   "","",rowArray[4],rowArray[5]]
                             """
                             #临时调整模板
                             #"代号",	"内容",	"类型",	"值域","数据处理方法","单位","区间"
                             busdata.loc[index] = [rowArray[msgMapIndex.loc["内容"]],"",bitType,newType,
                                                   "","",rowArray[msgMapIndex.loc["单位"]],\
                                                   rowArray[msgMapIndex.loc["备注"]]]
                             
                             "仅数据的第一行赋值信息名称等信息，其他为空,使msgarray与busdata对齐"
                             "如果为第一行数据则msgArray不赋值,另信息名称与第一行数据对齐"
                             if dataStartIndex == iterNum:
                                 pass
                             else:
                                 msgNameArray.append("")
                                 msgAddrArray.append("")
                                 msgIDArray.append("")
                             #只有成功赋值了才可以增长
                             index += 1
                     iterNum += 1
            else:
                print( "get_bus_msg:",
                      "没有找到正确的表头提取数据")
        else:
            pass
                
    busdataset = DataSet()
    busdataset.busdata = busdata
    busdataset.msgAddrArray = msgAddrArray
    busdataset.msgIDArray = msgIDArray
    busdataset.msgNameArray = msgNameArray
    resize_dataset(busdataset.msgNameArray, busdataset.msgAddrArray, 
                  busdataset.msgIDArray, busdataset.busdata,filename,outputpath)

def resize_dataset(msgName,msgAddr,msgId,data,filename,outputpath):
    "把数据整理为csv中的排列格式"
    "通过datanum将msg维数与data维数对齐"
    "校验消息维数"
    dataShape = data.shape
    emptyArray = np.zeros(dataShape[0],dtype = str)    
    if dataShape[0] == len(msgName):
        pass
    else:
        print("resize_dataset dimension error")
    msgDim = len(msgAddr)
    xyCode = list(emptyArray)
    xsCode = list(emptyArray)
    subAddrCode =list(emptyArray)
    lengthCode = list(emptyArray)
    i = 0
    while i < msgDim:
        if msgAddr[i] == "":
            pass
        else:
            "解析地址等"
            (xy,xs,subaddr,length) = ArrayCmp.get_bus_msg_param(msgAddr[i])
            
            xyCode[i] = xy
            xsCode[i] = xs
            subAddrCode[i] = subaddr
            lengthCode[i] = length
            print(i,":",msgAddr[i])
        i += 1
      
 
    tmpData = {"名称":msgName,"信源系统码":emptyArray,"信源机器码":xyCode,
               "信宿系统码":emptyArray,"信宿机器码":xsCode,
               "子地址":subAddrCode,"数据段长度":lengthCode,"ID":emptyArray}
    msgData = pd.DataFrame(tmpData)
    csvData = pd.concat([msgData,data],axis = 1)
    WriteToExcel(csvData, filename,outputpath)

def WriteToExcel(df,filename,outputpath):
    """
    获取dataframe，sheet名字为name
    一个datatframe对应一个excel
    excel的名称=name
    Parameters
    ----------
    df : TYPE
        DESCRIPTION.
    name : TYPE
        DESCRIPTION.

    Returns
    -------
    None.

    """
    "worksheet名称不能超过21个char"
    name = ArrayCmp.transfer_shortname(filename)       
        
    csvFileName = outputpath + name + '.xlsx'
    writer = pd.ExcelWriter(csvFileName)
    df.to_excel(writer,sheet_name = name,index = False)
    writer.save()
    writer.close()
    
    
def get_Eth_msg(doc,filename,outputpath):
    """
    解析网络通信协议，返回dataframe

    Parameters
    ----------
    doc : TYPE
        文档句柄.

    Returns
    -------
    包含信息的dataset类

    """
    "先把msg头部的各信息存到一个list里面，然后再重建维数一样的dataframe"
    templateName = ["名称",	"信源系统码",	"信源机器码",	"信宿系统码",\
                    "信宿机器码",	"子地址或消息地址","数据段长度（总线为字，其他为字节）",\
                    "ID","内容",	"子内容",	"类型（bit）",	"转换类型"	,"判读公式（暂不设计）",	\
                        "转换公式（变量必须为x）",	"单位",	"备注"]
    portBaseLine = ["序号",	"信源",	"信宿",	"信息内容",
                   "接收组播地址",	"接收端口号",	"信源系统码",	
                   "信源机器码",	"信宿系统码"	,"信宿机器码"]
    portApIndexArray = range(3,len(portBaseLine))
    startTableIndex = 1
    portDataStartIndex = 1
    portHeadIndex = 0
    (portData,startTableIndex) = get_appoint_table_content(doc, portBaseLine, 
                                         portApIndexArray, startTableIndex, 
                                         portDataStartIndex,portHeadIndex)
    "调换属性顺序与模板一致"
    msgDataset = portData[["信息内容","信源系统码",	
                   "信源机器码",	"信宿系统码"	,"信宿机器码", "接收组播地址"]]
    "row col"
    msgDataShape = msgDataset.shape
    "获取id"
    idBaseLine = ["序号","信源","信宿","信息内容","消息ID"]
    "startTableIndex 用上一个返回值"
    idApIndexArray = range(3,len(idBaseLine))
    idDataStartIndex = 1
    portHeadIndex = 0
    (idTmpData,startTableIndex)= get_appoint_table_content(doc, idBaseLine, 
                                         idApIndexArray, startTableIndex,
                                         idDataStartIndex, portHeadIndex)
    
    "添加一列空行，网络数据数据段长度不填"
    emptyArray = np.zeros(msgDataShape[0],dtype = str) 
    msgDataset.insert(msgDataShape[1],"数据段长度（总线为字，其他为字节）",emptyArray)
    "根据端口表的信息内容顺序调换id的,考虑索引找不到的特殊处理"
    idData = ArrayCmp.sort_dataframe_with_content(idTmpData, '信息内容',portData["信息内容"].values.tolist())
    msgDataset.insert(msgDataShape[1]+1,"ID",idData[["消息ID"]])
    "修改表格的列名与模板一致"
    msgDataset = msgDataset.rename(columns = {portBaseLine[3]:templateName[0],portBaseLine[4]:templateName[5]})

    "读取数据段，并对根据匹配关系获取消息其他信息，创建一个消息组"
    dataBaseLine = ["序号","参数","数据类型","数据长度（字节）","值域","单位","备注"]
    "先只要参数、数据类型、单位、备注，值域需进一步处理"
    "参数，数据类型，单位，备注"
    dataApIndexArray = [1,2,5,6]
    msgNum = msgDataShape[1]
 
    "创建无类型bit等列的数组"
    tpNamePart = msgDataset.columns.values.tolist() + ["参数","数据类型","单位","备注"]
    dataAllset = pd.DataFrame(columns = tpNamePart)
    "根据消息数量遍历表会有冗余"
    indexMsg = 0
    colNum = len(msgDataset.columns)
    namelist = msgDataset[["名称"]].values.tolist()
    while indexMsg < msgNum:
        
        "遍历接口数据约束部分的表格，总数为消息名称的维数"
        "startTableIndex每次在函数体内递增"
        (ethMsgData, msgNameTmp, msgPNameTmp, startTableIndex)\
            = get_appoint_table_content_multi(doc, dataBaseLine, \
                                              dataApIndexArray,startTableIndex, 2, 1)
        "检查信息头不是否有信息名，如果没有不解析"
        if "" == msgNameTmp or [msgNameTmp] not in namelist:
            pass
        else:
            "处理类型、添加判读、转换公式空行"
            
            "把数据和消息信息对齐"
            "先获取有多少行"
            rowNum = len(ethMsgData.index)
            emptyset = np.zeros([rowNum,colNum],dtype = str)
            "第一行赋值有效值"
            withNameEmptySet = pd.DataFrame(emptyset,columns=msgDataset.columns.values.tolist())
            "需要根据名字定为行数,关键词外面必须套[]"
            namelocindex = namelist.index([msgNameTmp])
            withNameEmptySet.loc[0,:] = msgDataset.loc[namelocindex,:]
            "链接两个数组"
            tmpAllset = pd.concat([withNameEmptySet,ethMsgData],axis = 1)
            dataAllset = dataAllset.append(tmpAllset,ignore_index=True)
            
        indexMsg += 1
        
    "对数据类型和其他空列进行处理"
    dataAllsetComplete = TransferTypeAndAddEmpty(dataAllset,templateName) 
    "写到csv中"
    WriteToExcel(dataAllsetComplete, filename,outputpath)
def get_uart_msg(doc,filename,outputpath):
    """
    读取串行通信协议数据

    Parameters
    ----------
    doc : TYPE
        DESCRIPTION.
    filename : TYPE
        名字用于csv文件名称和sheet名称.

    Returns
    -------
    None.

    """
    templateName = ["名称",	"信源系统码",	"信源机器码",	"信宿系统码",\
                    "信宿机器码",	"子地址或消息地址","数据段长度（总线为字，其他为字节）",\
                    "ID","内容",	"子内容",	"类型（bit）",	"转换类型"	,"判读公式（暂不设计）",	\
                        "转换公式（变量必须为x）",	"单位",	"备注"]
 
    startTableIndex = 0
    "获取id"
    idBaseLine = ["ID序号","ID定义","是否有数据"]
    "startTableIndex 用上一个返回值"
    idApIndexArray = [0,1]
    idDataStartIndex = 1
    portHeadIndex = 0
    (idTmpData,startTableIndex)= get_appoint_table_content(doc, idBaseLine, 
                                         idApIndexArray, startTableIndex,
                                         idDataStartIndex, portHeadIndex)
    msgDataset = pd.DataFrame(index = idTmpData.index,columns=[templateName[0],templateName[7]])
    msgDataset.loc[:,templateName[0]] = idTmpData.loc[:,idTmpData.columns[1]]
    msgDataset.loc[:,templateName[7]] = idTmpData.loc[:,idTmpData.columns[0]]
    
    

    "读取数据段，并对根据匹配关系获取消息其他信息，创建一个消息组"
    dataBaseLine = ["序号","参数","数据类型","数据长度（字节）","值域","单位","备注"]
    "先只要参数、数据类型、单位、备注，值域需进一步处理"
    "参数，数据类型，单位，备注"
    dataApIndexArray = [1,2,5,6]
    "设置为剩余表格数"
    msgNum = len(doc.tables) - startTableIndex + 1
 
    "创建无类型bit等列的数组"
    tpNamePart = msgDataset.columns.values.tolist() + ["参数","数据类型","单位","备注"]
    dataAllset = pd.DataFrame(columns = tpNamePart)
    "根据消息数量遍历表会有冗余"
    indexMsg = 0
    colNum = len(msgDataset.columns)
    namelist = msgDataset[["名称"]].values.tolist()
    while indexMsg < msgNum:
        
        "遍历接口数据约束部分的表格，总数为消息名称的维数"
        "startTableIndex每次在函数体内递增"
        (ethMsgData, msgNameTmp, msgPNameTmp, startTableIndex)\
            = get_appoint_table_content_multi(doc, dataBaseLine, \
                                              dataApIndexArray,startTableIndex, 2, 1)
        "检查信息头不是否有信息名，如果没有不解析"
        """后续加入子内容，可先判是否信息名称为空如果不空，再判断不在列表，
        如果不在列表检索是否在已有的字段中，如果在则更换二级表头为bit表头，再次读取表格
        根据读取内容切分表格，再concat
        """
        
        if "" == msgNameTmp or [msgNameTmp] not in namelist:
            pass
        else:
            "处理类型、添加判读、转换公式空行"
            
            "把数据和消息信息对齐"
            "先获取有多少行"
            rowNum = len(ethMsgData.index)
            emptyset = np.zeros([rowNum,colNum],dtype = str)
            "第一行赋值有效值"
            withNameEmptySet = pd.DataFrame(emptyset,columns=msgDataset.columns.values.tolist())
            "需要根据名字定为行数,关键词外面必须套[]"
            namelocindex = namelist.index([msgNameTmp])
            withNameEmptySet.loc[0,:] = msgDataset.loc[namelocindex,:]
            "链接两个数组"
            tmpAllset = pd.concat([withNameEmptySet,ethMsgData],axis = 1)
            dataAllset = dataAllset.append(tmpAllset,ignore_index=True)
            
        indexMsg += 1
        
    "对数据类型和其他空列进行处理"
    dataAllsetComplete = TransferTypeAndAddEmpty(dataAllset,templateName) 
    "写到csv中"
    WriteToExcel(dataAllsetComplete, filename,outputpath)
def TransferTypeAndAddEmpty(df,tpName):
    """
    从协议中获取内容、数据类型、单位、备注后
    转换数据类型，添加判读公式和转换公式（空行）
    属性与模板修改一致
    Parameters
    ----------
    df : TYPE
        "参数，数据类型，单位，备注".

    Returns
    -------
    None.

    """
    "添加数据类型"
    #此处后面优化，确定两个列含义相似
    "换名字"
    df = df.rename(columns={'数据类型' : '转换类型','参数':'内容'})
    dataShape = df.shape
    emptyArray = np.zeros(dataShape[0],dtype = str)  
    dataType = df.loc[:,"转换类型"]
    bitDataType = np.zeros(len(dataType))
    newDataType = pd.Series(index = df.index,dtype = str)

    for index,tmpstr in enumerate(dataType) :
        try:
            (bitDataType[index],newDataType.loc[index]) = ArrayCmp.standard_datatype(tmpstr)
            
        except ValueError:
            print("TransferTypeAndAddEmpty:",tmpstr)
            
        
    "查找数据类型索引"
    dfHeadName = df.columns.values.tolist()
    cp_list = ArrayCmp.compare_list_return_index(tpName, dfHeadName)
    addNum = len(cp_list)
    i = 0
    while i < addNum:
        df.insert((cp_list[i] ),tpName[cp_list[i]],emptyArray)
        i += 1
    
    "对类型赋值"
    df.loc[:,tpName[10]] = pd.Series(bitDataType)
    "转换类型赋值"
    df.loc[:,tpName[11]] = pd.Series(newDataType)
    return df
 
        
    
            
    
def get_appoint_table_content(doc,baseline,appointIndexArray, 
                              startTableIndex,dataStartIndex,headIndex):
    """
    从指定表格索引开始找到表头与baselin匹配的表格，并获取内容

    Parameters
    ----------
    doc : TYPE
        DESCRIPTION.
    baseline : 元组
        表头
    appointIndexArray : 元组
        用于控制获取哪些列
    startTableIndex : TYPE
        控制从哪个table开始找下一个table
    dataStartIndex：
        数据从第几行开始
    headIndex:
        消息头从第几行开始

    Returns
    -------
    dataframe

    """
    """匹配总线信息表头，如果是对的表头则提取数据放入dataframe"""
    "遍历文档表格，直至找到第一个符合表头要求的表格"
    
    """busbaseline = ["信息名称","信息标识"]"""
    """msgbaseline = ["序号",	"信源",	"信宿",	"信息内容",
                   "接收组播地址",	"接收端口号",	"信源系统码",	
                   "信源机器码",	"信宿系统码"	,"信宿机器码"]
    """
    

    rowArray = []
    "数据字段的表头索引，防止存在列表合并单元导致取错数据"
    indexArray = []
    indexNameArray = []
    ArrayCmp.extract_appoint_list(baseline, appointIndexArray, indexNameArray)
    busdata = pd.DataFrame(columns = indexNameArray)
    index = 0
    tableIndex = startTableIndex
    tablenum = len(doc.tables)
    while tableIndex < tablenum:
        "匹配表头"
        tmpTable = doc.tables[tableIndex]
        get_doc_table_row_context(tmpTable, rowArray, headIndex,False)
        get_doc_table_head_index(tmpTable, rowArray, headIndex,indexArray)
        if mate_table_head(rowArray, baseline):
            "提取数据"
            tableRowNum = len(tmpTable.rows)
            iterNum = dataStartIndex

            while iterNum < tableRowNum:   
                 "取内容,类型当全部写完后批量转换,防止合并单元格影响取数采用表头的索引"
                 get_doc_tabel_row_appoint_context(tmpTable,rowArray, iterNum, indexArray)
                
                 "去掉注"
                 "如果不小于0则是有注"
                 if rowArray[0].find('注') >= 0:
                     pass
                 else:
                     if len(rowArray) != len(baseline):
                         print("get_appoint_table_content:表格数据列错误")
                     else:
                         busdata.loc[index] = ArrayCmp.get_appoint_array(rowArray, appointIndexArray)
                         "ArrayCmp.list_copy(tmplist, busdata.loc[index])"
                     
                     index += 1
                 iterNum += 1
             
            "前面的表已经遍历过了，再进入该函数可以从这个索引之后找"
            
            startTableIndex = tableIndex + 1
            break
        else:
            pass
        tableIndex += 1
                    
    return busdata,startTableIndex

def get_appoint_table_content_multi(doc,msgNameBaseLine,baseline,indexCtrlArray,appointIndexArray):
                                    
    """
    适用于多层嵌套表格，从指定表格索引开始找到表头与baselin匹配的表格，并获取内容
    只要找到一个就推出

    Parameters
    ----------
    doc : TYPE
        DESCRIPTION.
    msgNameBaseLine：list
        描述消息名称的表头    
    baseline : 元组
        描述字段属性表头
    appointIndexArray : 元组
        用于控制获取哪些列
    indexCtrlArray:list
    内容分为：
    startTableIndex : TYPE
        控制从哪个table开始找下一个table
    msgNameIndex:
        描述消息名称的行从第几行开始
    headIndex:
    数据消息头从第几行开始
    dataStartIndex：
        数据从第几行开始


    Returns
    -------
    dataframe
    msgName:该消息的名称
    msgParentName:上级消息的名称

    """
    """匹配总线信息表头，如果是对的表头则提取数据放入dataframe"""
    "遍历文档表格，直至找到第一个符合表头要求的表格"
    #先检查indexCtrlArray维数和数值合法性
    if len(indexCtrlArray) != 4:
        print("get_appoint_table_content_multi:indexCtrlArray dimension error")
        return
    else:    

        if indexCtrlArray[1] >= indexCtrlArray[2] or indexCtrlArray[2] >= indexCtrlArray[3]:
            print("get_appoint_table_content_multi:indexCtrlArray value error")
            return
        else:
            startTableIndex = indexCtrlArray[0]
            msgNameIndex = indexCtrlArray[1]
            headIndex = indexCtrlArray[2]
            dataStartIndex = indexCtrlArray[3]
    
    #检查msgNameBaseLine维数
    if len(msgNameBaseLine) != 2:
        print("get_appoint_table_content_multi:msgNameBaseLine dimension error")
        return
    
    
    rowArray = []
    rowOldArray = []
    "数据字段的表头索引，防止存在列表合并单元导致取错数据"
    indexArray = []
    indexNameArray = []
    ArrayCmp.extract_appoint_list(baseline, appointIndexArray, indexNameArray)
    busdata = pd.DataFrame(columns = indexNameArray)
    index = 0
    tableIndex = startTableIndex
    tablenum = len(doc.tables)
    msgName = ""
    msgParentName = ""
    while tableIndex < tablenum:
        "匹配表头"
        tmpTable = doc.tables[tableIndex]
        prevent_table_break(tmpTable)
        #读到指定的消息头行数      
        "访问表头，获取信息名称，默认信息名称表头在第一行"
        get_doc_table_row_context(tmpTable, rowOldArray, msgNameIndex,False)
        #把符号去掉
        ArrayCmp.get_clean_chsArray(rowOldArray, rowArray)

        if len(rowArray) > 3:
            if rowArray[0] == msgNameBaseLine[0] and rowArray[2] == msgNameBaseLine[1]:
                msgName = rowArray[1]
                msgParentName = rowArray[3]

                get_doc_table_row_context(tmpTable, rowOldArray, headIndex,False)
                get_doc_table_head_index(tmpTable, rowOldArray, headIndex,indexArray)
                ArrayCmp.get_clean_chsArray(rowOldArray, rowArray)
                if mate_table_head(rowArray, baseline):
                    "提取数据"
                    tableRowNum = len(tmpTable.rows)
                    iterNum = dataStartIndex
        
                    while iterNum < tableRowNum:   
                         "取内容,类型当全部写完后批量转换,防止合并单元格影响取数采用表头的索引"
                         get_doc_tabel_row_appoint_context(tmpTable,rowArray, iterNum, indexArray)
                        
                         "去掉注"
                         "如果不小于0则是有注"
                         if rowArray[0].find('注') >= 0:
                             pass
                         else:
                             if len(rowArray) != len(baseline):
                                 print("get_appoint_table_content_multi:表格数据列错误")
                             else:
                                 busdata.loc[index] =  ArrayCmp.get_appoint_array(rowArray, appointIndexArray)
                             
                             index += 1
                         iterNum += 1
                     
                    #更新下次遍历的索引
                    tableIndex += 1
                    "前面的表已经遍历过了，再进入该函数可以从这个索引之后找"
                    indexCtrlArray[0] = tableIndex
                    break
                else:
                    print("get_appoint_table_contnet_multi:第二行信息头错误文档中表头为：")
                    print(rowArray)
            else:
                print("get_appoint_table_contnet_multi: 第一行信息头错误:",rowArray)
        else:
            print("get_appoint_table_contnet_multi: 第一行信息头不满足四列",rowArray)
            
        tableIndex += 1
        print("tableindex = ", tableIndex)
    #如果都遍历结束了也没有返回           
    indexCtrlArray[0] = tableIndex
                    
    return busdata,msgName,msgParentName,indexCtrlArray

def get_msg_data(doc,filename,outputpath,msgNameBaseLine = None,dataBaseLine = None,indexCtrlArray = None,dataApIndexArray = None):
    """
    仅提取信息流数表，不关联ID地址等信息流信息

    Parameters
    ----------
    doc : TYPE
        DESCRIPTION.
    filename : TYPE
        DESCRIPTION.
    outputpath : TYPE
        DESCRIPTION.
    dataBaseLine:list
        表头名称
    dataApIndexArray:list
    "参数，数据类型，单位，备注"的索引值，顺序与要求保持一致
    Returns
    -------
    None.

    """
    templateName = ["名称",	"信源系统码",	"信源机器码",	"信宿系统码",\
                    "信宿机器码",	"子地址或消息地址","数据段长度（总线为字，其他为字节）",\
                    "ID","内容",	"子内容",	"类型（bit）",	"转换类型"	,"判读公式（暂不设计）",	\
                        "转换公式（变量必须为x）",	"单位",	"备注"]
 
    
    #后面要设置index，名称，ID
    msgDataset = pd.DataFrame(columns=[templateName[0],templateName[7]])
    
    if msgNameBaseLine == None:
        msgNameBaseLine = ["信息名称","上级信息名称"]
    else:
        print('msgNameBaseline = ',msgNameBaseLine)

    "读取数据段，并对根据匹配关系获取消息其他信息，创建一个消息组"
    if dataBaseLine == None:
        dataBaseLine = ["序号","参数","数据类型","数据长度（字节）","值域","单位","备注"]
    else:
        print('dataBaseline is ',dataBaseLine)
        
    if indexCtrlArray == None:
        indexCtrlArray = [0,0,1,2]
    else:
        print('indexCtrlArray = ',indexCtrlArray)
        
        
    "先只要参数、数据类型、单位、备注，值域需进一步处理"
    "参数，数据类型，单位，备注"
    if dataApIndexArray == None:
         dataApIndexArray = [1,2,5,6]
    else:
        print('get index name is ')
        for i in dataApIndexArray:
            print(dataBaseLine[i])
    #最小为0
    startTableIndex = 0    
         
    "设置为剩余表格数"
    msgNum = len(doc.tables) - startTableIndex
 
    "创建无类型bit等列的数组"
    tpNamePart = msgDataset.columns.values.tolist() + ["参数","数据类型","单位","备注"]
    dataAllset = pd.DataFrame(columns = tpNamePart)
    "根据消息数量遍历表会有冗余"
    indexMsg = 0
    colNum = len(msgDataset.columns)
 
    while indexMsg < msgNum:
        
        "遍历接口数据约束部分的表格，总数为消息名称的维数"
        "startTableIndex每次在函数体内递增"
        (ethMsgData, msgNameTmp, msgPNameTmp, indexCtrlArray)\
            = get_appoint_table_content_multi\
                (doc, msgNameBaseLine,dataBaseLine, indexCtrlArray, dataApIndexArray)
        print("get msg:",msgNameTmp)
        #重置ethMsgData的属性,否则影响concat
        ethMsgData.columns = ["参数","数据类型","单位","备注"]
        "检查信息头不是否有信息名，如果没有不解析"
        """后续加入子内容，可先判是否信息名称为空如果不空，再判断不在列表，
        如果不在列表检索是否在已有的字段中，如果在则更换二级表头为bit表头，再次读取表格
        根据读取内容切分表格，再concat
        """
        if msgNameTmp == "":
            indexMsg += 1
            pass
        else:
        
            "处理类型、添加判读、转换公式空行"
            
            "把数据和消息信息对齐"
            "先获取有多少行"
            rowNum = len(ethMsgData.index)
            emptyset = np.zeros([rowNum,colNum],dtype = str)
            "第一行赋值有效值"
            withNameEmptySet = pd.DataFrame(emptyset,columns=msgDataset.columns.values.tolist())
            #取信息头中信息名称     
            withNameEmptySet.loc[0,'名称'] = msgNameTmp
            "链接两个数组"
            tmpAllset = pd.concat([withNameEmptySet,ethMsgData],axis = 1)
            dataAllset = dataAllset.append(tmpAllset,ignore_index=True)
        #用返回值指向下次遍历的值
        if indexCtrlArray[0] <= indexMsg:
            print("error:回退")
        indexMsg = indexCtrlArray[0]
 
    "对数据类型和其他空列进行处理"
    dataAllsetComplete = TransferTypeAndAddEmpty(dataAllset,templateName) 
    "写到csv中"
    WriteToExcel(dataAllsetComplete, filename,outputpath)
    
      
def get_protocol_info(fileAbsPath,filename = None):
    """
    读取协议，
    如果为doc则转换为docx,
    根据协议名称进行类型转换生成csv

    Parameters
    ----------
    fileAbsPath : TYPE
        绝对路径.
    filename : TYPE
        文件名包含后缀,如果为空则遍历fileabspath

    Returns
    -------
    None.

    """
    "先进行docx批量转换"
   
    outputFilePath = fileAbsPath + "/output/"
    fm.transfer_protocolfile_type(fileAbsPath, outputFilePath)
        
    "遍历output文件夹进行解析"
    for docname in os.listdir(outputFilePath):
        docxname = outputFilePath + "/" + docname
        doc = Document(docxname)
        "提取表格名称"
        splilist = docname.split(".")
        shortname = splilist[0]
        "根据docname区分是哪一种协议"
        pType = ArrayCmp.extract_proto_type(shortname)
        if 0 == pType:
            "bus"
            get_bus_msg(doc,shortname,outputFilePath)
        elif 1 == pType:
            "uart"
            get_uart_msg(doc,shortname,outputFilePath)
        elif 2 == pType:
            "eth"
            get_Eth_msg(doc, shortname,outputFilePath)
        else:
            "三种类型均检索"
            mixname = "bus" +docname  
            get_bus_msg(doc,mixname,outputFilePath)
            mixname = "uart" +docname
            get_uart_msg(doc, mixname,outputFilePath)
            mixname = "eth" + docname
            get_Eth_msg(doc, mixname,outputFilePath)
                
def prevent_table_break(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblpPr = OxmlElement("w:tblpPr")
    tblpPr.set(qn("w:keepNext"),"true")
    tblPr.append(tblpPr)          
        
        
        